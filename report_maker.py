import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
import re
from deep_translator import GoogleTranslator
import language_tool_python
import time

translator = GoogleTranslator(source='es', target='en')

tool = None

def init_language_tool():
    global tool
    if tool is None:
        try:
            tool = language_tool_python.LanguageTool('en-US')
        except:
            pass
    return tool

import time

def translate_to_english(text):
    """
    Traduce texto de español a inglés con sistema de reintentos.
    Optimizado para uso individual con IPs únicas.
    """
    if not text or not text.strip():
        return text
    
    try:
        text = text.strip()
        MAX_CHUNK_SIZE = 250
        
        if len(text) <= MAX_CHUNK_SIZE:
            return translate_with_retry(text)
        
        sentences = text.replace('. ', '.|').replace('? ', '?|').replace('! ', '!|').split('|')
        translated_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            translated = translate_with_retry(sentence)
            translated_sentences.append(translated)
        
        return ' '.join(translated_sentences)
        
    except Exception as e:
        print(f"🚨 Error general: {e}")
        return text


def translate_with_retry(text, max_retries=3):
    """
    Intenta traducir con reintentos automáticos.
    Valida que la traducción no esté truncada.
    """
    original_length = len(text)
    
    for attempt in range(max_retries):
        try:
            translated = translator.translate(text)
            
            if len(translated) >= original_length * 0.3:
                return translated
            else:
                print(f"⚠️ Intento {attempt + 1}: Traducción corta")
                if attempt < max_retries - 1:
                    time.sleep(0.5)
                    
        except Exception as e:
            print(f"⚠️ Intento {attempt + 1} falló: {e}")
            if attempt < max_retries - 1:
                time.sleep(0.5)
    
    print(f"⚠️ No se pudo traducir después de {max_retries} intentos")
    return text


def translate_by_sentences(text):
    """Divide texto en oraciones y traduce cada una."""
    import time
    
    sentences = text.replace('. ', '.|').replace('? ', '?|').replace('! ', '!|').split('|')
    translated_sentences = []
    
    for i, sentence in enumerate(sentences):
        sentence = sentence.strip()
        if not sentence:
            continue
        
        try:
            if i > 0:
                time.sleep(0.2)
            
            translated = translator.translate(sentence)
            
            if len(translated) >= len(sentence) * 0.3:
                translated_sentences.append(translated)
            else:
                print(f"⚠️ Oración truncada, usando original")
                translated_sentences.append(sentence)
                
        except Exception as e:
            print(f"⚠️ Error en oración: {e}")
            translated_sentences.append(sentence)
    
    return ' '.join(translated_sentences)

def correct_grammar(text):
    if not text or not text.strip():
        return text
    try:
        lt = init_language_tool()
        if lt:
            matches = lt.check(text)
            return language_tool_python.utils.correct(text, matches)
        return text
    except:
        return text

def translate_and_correct(text):
    if not text or not text.strip():
        return text
    return correct_grammar(translate_to_english(text))

def translate_equipment_info(text):
    """
    Traduce información de equipo línea por línea.
    Usa división por oraciones para valores largos.
    """
    if not text or not text.strip():
        return text
    
    translations = {
        'nombre del equipo': 'Equipment', 'nombre': 'Equipment',
        'modelo': 'Model', 'número de serie': 'Serial Number', 'serial': 'Serial Number',
        'versión hardware': 'Hardware Version', 'versión software': 'Software Version',
        'versión firmware': 'Firmware Version', 'código de país': 'Country Code',
        'product id': 'Product ID', 'estado': 'State', 'versión': 'Version',
    }
    
    lines = text.split('\n')
    result = []
    
    for line in lines:
        line = line.strip()
        if not line:
            result.append('')
            continue
        
        if ':' in line:
            parts = line.split(':', 1)
            if len(parts) == 2:
                field = parts[0].strip().lower()
                value = parts[1].strip()
                
                if field in translations:
                    field_translated = translations[field]
                else:
                    try:
                        field_translated = translate_with_retry(parts[0].strip())
                    except:
                        field_translated = parts[0].strip()
                
                if len(value) > 80:
                    try:
                        value_translated = translate_to_english(value)
                    except:
                        value_translated = value
                else:
                    if any(word in value.lower() for word in ['configurado', 'activa', 'funcionando', 'habilitado', 'deshabilitado']):
                        try:
                            value_translated = translate_with_retry(value)
                        except:
                            value_translated = value
                    else:
                        value_translated = value
                
                result.append(f"{field_translated}: {value_translated}")
            else:
                result.append(line)
        else:
            try:
                line_translated = translate_to_english(line)
                result.append(line_translated)
            except:
                result.append(line)
    
    return '\n'.join(result)

class MaterialColors:
    PRIMARY = '#0078D4'
    PRIMARY_HOVER = '#005A9E'
    PRIMARY_LIGHT = '#4A9EDE'
    PRIMARY_DARK = '#004578'
    
    SUCCESS = '#107C10'
    SUCCESS_HOVER = '#0E6B0E'
    
    BG_LIGHT = '#F5F5F5'
    BG_CARD = '#FFFFFF'
    TEXT_PRIMARY = '#1A1A1A'
    TEXT_SECONDARY = '#707070'
    BORDER_LIGHT = '#E0E0E0'
    SHADOW = '#CCCCCC'
    
    SCROLLBAR_BG = '#F0F0F0'
    SCROLLBAR_ACTIVE = "#868686"

class ModernScrollbar(tk.Canvas):
    def __init__(self, parent, orient='vertical', command=None, **kwargs):
        width = 14 if orient == 'vertical' else 200
        height = 200 if orient == 'vertical' else 14
        
        super().__init__(parent, width=width, height=height,
                        bg=MaterialColors.SCROLLBAR_BG, highlightthickness=0, **kwargs)
        
        self.orient = orient
        self.command = command
        self.pressed = False
        self.thumb_pos = 0
        self.thumb_size = 0.3
        self.hover = False
        
        self.bind('<Button-1>', self.on_press)
        self.bind('<B1-Motion>', self.on_drag)
        self.bind('<ButtonRelease-1>', self.on_release)
        self.bind('<Enter>', self.on_enter)
        self.bind('<Leave>', self.on_leave)
        self.bind('<Configure>', self.on_configure)
        
        self.draw_thumb()
    
    def on_configure(self, event):
        self.draw_thumb()
    
    def draw_thumb(self):
        self.delete('all')
        
        if self.thumb_size <= 0 or self.thumb_size >= 1:
            return
        
        color = MaterialColors.SCROLLBAR_ACTIVE if (self.pressed or self.hover) else MaterialColors.TEXT_SECONDARY
        
        if self.orient == 'vertical':
            w = int(self['width'])
            h = self.winfo_height() if self.winfo_height() > 1 else 200
            
            thumb_height = max(30, int(h * self.thumb_size))
            thumb_y = int((h - thumb_height) * self.thumb_pos / (1 - self.thumb_size)) if self.thumb_size < 1 else 0
            
            x1, y1 = 3, thumb_y
            x2, y2 = w - 3, thumb_y + thumb_height
            
            self.create_rounded_rect(x1, y1, x2, y2, 4, fill=color, outline='')
        else:
            w = self.winfo_width() if self.winfo_width() > 1 else 200
            h = int(self['height'])
            
            thumb_width = max(30, int(w * self.thumb_size))
            thumb_x = int((w - thumb_width) * self.thumb_pos / (1 - self.thumb_size)) if self.thumb_size < 1 else 0
            
            x1, y1 = thumb_x, 3
            x2, y2 = thumb_x + thumb_width, h - 3
            
            self.create_rounded_rect(x1, y1, x2, y2, 4, fill=color, outline='')
    
    def create_rounded_rect(self, x1, y1, x2, y2, radius, **kwargs):
        points = [
            x1+radius, y1, x2-radius, y1, x2, y1, x2, y1+radius,
            x2, y2-radius, x2, y2, x2-radius, y2, x1+radius, y2,
            x1, y2, x1, y2-radius, x1, y1+radius, x1, y1
        ]
        return self.create_polygon(points, smooth=True, **kwargs)
    
    def set(self, first, last):
        first = float(first)
        last = float(last)
        
        self.thumb_pos = first
        self.thumb_size = last - first
        
        self.draw_thumb()
    
    def on_press(self, event):
        self.pressed = True
        
        if self.orient == 'vertical':
            h = self.winfo_height()
            thumb_height = max(30, int(h * self.thumb_size))
            thumb_y = int((h - thumb_height) * self.thumb_pos / (1 - self.thumb_size)) if self.thumb_size < 1 else 0
            
            if thumb_y <= event.y <= thumb_y + thumb_height:
                self.drag_start_y = event.y - thumb_y
            else:
                ratio = event.y / h
                if self.command:
                    self.command('moveto', ratio)
                self.drag_start_y = thumb_height / 2
        else:
            w = self.winfo_width()
            thumb_width = max(30, int(w * self.thumb_size))
            thumb_x = int((w - thumb_width) * self.thumb_pos / (1 - self.thumb_size)) if self.thumb_size < 1 else 0
            
            if thumb_x <= event.x <= thumb_x + thumb_width:
                self.drag_start_x = event.x - thumb_x
            else:
                ratio = event.x / w
                if self.command:
                    self.command('moveto', ratio)
                self.drag_start_x = thumb_width / 2
        
        self.draw_thumb()
    
    def on_drag(self, event):
        if not self.command:
            return
        
        if self.orient == 'vertical':
            h = self.winfo_height()
            thumb_height = max(30, int(h * self.thumb_size))
            max_y = h - thumb_height
            
            if max_y > 0:
                new_y = event.y - self.drag_start_y
                ratio = new_y / max_y
                ratio = max(0, min(ratio, 1))
                self.command('moveto', ratio * (1 - self.thumb_size))
        else:
            w = self.winfo_width()
            thumb_width = max(30, int(w * self.thumb_size))
            max_x = w - thumb_width
            
            if max_x > 0:
                new_x = event.x - self.drag_start_x
                ratio = new_x / max_x
                ratio = max(0, min(ratio, 1))
                self.command('moveto', ratio * (1 - self.thumb_size))
    
    def on_release(self, event):
        self.pressed = False
        self.draw_thumb()
    
    def on_enter(self, event):
        self.hover = True
        self.draw_thumb()
    
    def on_leave(self, event):
        self.hover = False
        if not self.pressed:
            self.draw_thumb()

class RoundedButton(tk.Canvas):
    def __init__(self, parent, text, command, bg_color, fg_color='white', 
                 hover_color=None, font=('Segoe UI', 11, 'bold'), 
                 width=120, height=40, corner_radius=6):
        super().__init__(parent, width=width, height=height, 
                        bg=parent['bg'], highlightthickness=0)
        
        self.command = command
        self.bg_color = bg_color
        self.fg_color = fg_color
        self.hover_color = hover_color or bg_color
        self.font = font
        self.corner_radius = corner_radius
        self.width = width
        self.height = height
        self.text = text
        
        self.draw_button(self.bg_color)
        
        self.bind('<Button-1>', self.on_click)
        self.bind('<Enter>', self.on_enter)
        self.bind('<Leave>', self.on_leave)
        
    def draw_button(self, color):
        self.delete('all')
        self.create_rounded_rect(2, 2, self.width-2, self.height-2, 
                                self.corner_radius, fill=color, outline='')
        self.create_text(self.width/2, self.height/2, text=self.text,
                        fill=self.fg_color, font=self.font)
    
    def create_rounded_rect(self, x1, y1, x2, y2, radius, **kwargs):
        points = [
            x1+radius, y1, x2-radius, y1, x2, y1, x2, y1+radius,
            x2, y2-radius, x2, y2, x2-radius, y2, x1+radius, y2,
            x1, y2, x1, y2-radius, x1, y1+radius, x1, y1
        ]
        return self.create_polygon(points, smooth=True, **kwargs)
    
    def on_click(self, event):
        if self.command:
            self.command()
    
    def on_enter(self, event):
        self.draw_button(self.hover_color)
        self.configure(cursor='hand2')
    
    def on_leave(self, event):
        self.draw_button(self.bg_color)
        self.configure(cursor='')

class AutoNumberedText(tk.Text):
    def __init__(self, master=None, **kwargs):
        super().__init__(master, **kwargs)
        self.bind('<Return>', self.auto_number)
        self.bind('<BackSpace>', self.handle_backspace)
        self.bind('<Delete>', self.handle_delete)
        self.configure(undo=True, maxundo=-1)
        self._renumbering = False
    
    def auto_number(self, event):
        # Obtener la posición actual del cursor
        current_line = self.index('insert').split('.')[0]
        current_line_num = int(current_line)
        
        # Obtener el contenido de la línea actual
        current_content = self.get(f"{current_line}.0", f"{current_line}.end")
        
        # Extraer el número actual si existe
        match = re.match(r'^(\d+)\.\s*', current_content)
        if match:
            next_num = int(match.group(1)) + 1
        else:
            # Si no hay número, contar las líneas con contenido
            content = self.get('1.0', 'end-1c')
            lines = [l.strip() for l in content.split('\n') if l.strip()]
            next_num = len(lines) + 1
        
        # Insertar nueva línea con numeración
        self.insert('insert', f'\n{next_num}. ')
        
        # Renumerar líneas posteriores
        self.after(1, lambda: self.renumber_from_line(current_line_num + 2, next_num + 1))
        
        return 'break'
    
    def handle_backspace(self, event):
        """Maneja el borrado con BackSpace y renumera inmediatamente"""
        if self._renumbering:
            return None
        
        try:
            # Obtener la selección actual
            if self.tag_ranges("sel"):
                # Si hay texto seleccionado, programar renumeración después del borrado
                self.after(1, self.renumber_all_lines)
                return None
            
            # Obtener posición actual
            current_pos = self.index('insert')
            current_line_num = int(current_pos.split('.')[0])
            current_col = int(current_pos.split('.')[1])
            
            # Si estamos al inicio de una línea (col 0)
            if current_col == 0 and current_line_num > 1:
                # Vamos a unir con la línea anterior
                prev_line = self.get(f"{current_line_num-1}.0", f"{current_line_num-1}.end")
                current_line = self.get(f"{current_line_num}.0", f"{current_line_num}.end")
                
                # Si la línea actual está numerada y tiene contenido
                match_current = re.match(r'^\d+\.\s*(.*)$', current_line)
                if match_current and match_current.group(1).strip():
                    # Programar renumeración después de la unión
                    self.after(1, self.renumber_all_lines)
            else:
                # Obtener contenido de la línea actual
                line_content = self.get(f"{current_line_num}.0", f"{current_line_num}.end")
                
                # Si después del backspace la línea quedará vacía o solo con número
                if len(line_content.strip()) <= 3:  # "X. " o menos
                    self.after(1, self.renumber_all_lines)
        except:
            pass
        
        return None
    
    def handle_delete(self, event):
        """Maneja el borrado con Delete y renumera inmediatamente"""
        if self._renumbering:
            return None
        
        try:
            # Si hay selección, programar renumeración
            if self.tag_ranges("sel"):
                self.after(1, self.renumber_all_lines)
                return None
            
            # Obtener posición actual
            current_pos = self.index('insert')
            current_line_num = int(current_pos.split('.')[0])
            line_content = self.get(f"{current_line_num}.0", f"{current_line_num}.end")
            current_col = int(current_pos.split('.')[1])
            
            # Si estamos al final de la línea, se unirá con la siguiente
            if current_col >= len(line_content):
                self.after(1, self.renumber_all_lines)
            else:
                # Si la línea quedará vacía después del delete
                if len(line_content.strip()) <= 3:
                    self.after(1, self.renumber_all_lines)
        except:
            pass
        
        return None
    
    def renumber_from_line(self, start_line, start_num):
        """Renumera las líneas desde start_line con numeración consecutiva"""
        if self._renumbering:
            return
        
        self._renumbering = True
        try:
            total_lines = int(self.index('end-1c').split('.')[0])
            current_num = start_num
            
            for line_num in range(start_line, total_lines + 1):
                line_content = self.get(f"{line_num}.0", f"{line_num}.end")
                
                # Si la línea tiene numeración, actualizarla
                match = re.match(r'^\d+\.\s*(.*)$', line_content)
                if match:
                    rest_of_line = match.group(1)
                    if rest_of_line.strip():  # Solo si hay contenido
                        self.delete(f"{line_num}.0", f"{line_num}.end")
                        self.insert(f"{line_num}.0", f"{current_num}. {rest_of_line}")
                        current_num += 1
        finally:
            self._renumbering = False
    
    def renumber_all_lines(self):
        """Renumera TODAS las líneas del documento manteniendo orden consecutivo"""
        if self._renumbering:
            return
        
        self._renumbering = True
        try:
            # Guardar posición del cursor
            try:
                cursor_pos = self.index('insert')
                cursor_line = int(cursor_pos.split('.')[0])
                cursor_col = int(cursor_pos.split('.')[1])
            except:
                cursor_pos = '1.0'
                cursor_line = 1
                cursor_col = 0
            
            content = self.get('1.0', 'end-1c')
            lines = content.split('\n')
            
            current_num = 1
            line_number_changes = {}  # Mapeo de número viejo a nuevo
            
            for line_idx, line in enumerate(lines):
                line_num = line_idx + 1
                
                # Si la línea tiene numeración
                match = re.match(r'^(\d+)\.\s*(.*)$', line)
                if match:
                    old_num = int(match.group(1))
                    rest_of_line = match.group(2).strip()
                    
                    if rest_of_line:  # Solo si hay contenido después del número
                        line_number_changes[old_num] = current_num
                        expected_text = f"{current_num}. {rest_of_line}"
                        current_text = self.get(f"{line_num}.0", f"{line_num}.end")
                        
                        # Solo actualizar si cambió
                        if current_text != expected_text:
                            self.delete(f"{line_num}.0", f"{line_num}.end")
                            self.insert(f"{line_num}.0", expected_text)
                        
                        current_num += 1
                    else:
                        # Línea con solo número, eliminarla
                        self.delete(f"{line_num}.0", f"{line_num}.end")
            
            # Restaurar posición del cursor de manera inteligente
            try:
                new_cursor_pos = f"{cursor_line}.{cursor_col}"
                self.mark_set('insert', new_cursor_pos)
            except:
                try:
                    self.mark_set('insert', '1.0')
                except:
                    pass
                    
        except Exception as e:
            pass
        finally:
            self._renumbering = False
    
    def get_numbered_text(self):
        content = self.get('1.0', 'end-1c')
        lines = content.split('\n')
        clean = []
        counter = 1
        for line in lines:
            clean_line = re.sub(r'^\d+\.\s*', '', line).strip()
            if clean_line:
                clean.append(f"{counter}. {clean_line}")
                counter += 1
        return '\n'.join(clean)
    
class RepairReportGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("ReportMaker v1.2.5")
        self.root.geometry("1450x850")
        self.root.configure(bg=MaterialColors.BG_LIGHT)
        self.root.minsize(1200, 700)
        
        self.setup_modern_style()
        
        self.summary_widgets = []
        self.procedure_widgets = []
        self.expected_widgets = []
        self.always_visible_widgets = []
        
        self.create_widgets()

    def setup_modern_style(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        style.configure('Modern.TCombobox',
                       fieldbackground='white',
                       background=MaterialColors.PRIMARY,
                       foreground=MaterialColors.TEXT_PRIMARY,
                       borderwidth=1,
                       relief='flat',
                       arrowsize=15)
        
        style.map('Modern.TCombobox',
                 fieldbackground=[('readonly', 'white')],
                 selectbackground=[('readonly', MaterialColors.PRIMARY)],
                 selectforeground=[('readonly', 'white')])
    
    def create_rounded_frame(self, parent, **kwargs):
        frame = tk.Frame(parent, **kwargs)
        frame.configure(highlightbackground=MaterialColors.BORDER_LIGHT,
                       highlightthickness=1,
                       relief=tk.FLAT)
        return frame
    
    def create_widgets(self):
        header = tk.Frame(self.root, bg=MaterialColors.PRIMARY, height=70)
        header.pack(fill=tk.X, side=tk.TOP)
        header.pack_propagate(False)
        
        shadow = tk.Frame(self.root, bg=MaterialColors.SHADOW, height=2)
        shadow.pack(fill=tk.X, side=tk.TOP)
        
        title_frame = tk.Frame(header, bg=MaterialColors.PRIMARY)
        title_frame.pack(pady=18)
        
        tk.Label(title_frame, text="ReportMaker", font=('Segoe UI', 20, 'bold'),
                bg=MaterialColors.PRIMARY, fg='white').pack(side=tk.LEFT)
        tk.Label(title_frame, text=" v1.2.5", font=('Segoe UI', 11),
                bg=MaterialColors.PRIMARY, fg='#CCCCCC').pack(side=tk.LEFT, padx=(5, 0))
        
        main_paned = tk.PanedWindow(self.root, orient=tk.HORIZONTAL, 
                                     bg=MaterialColors.BG_LIGHT, 
                                     sashwidth=8, sashrelief=tk.FLAT,
                                     bd=0)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        left_container = tk.Frame(main_paned, bg=MaterialColors.BG_LIGHT)
        
        canvas = tk.Canvas(left_container, bg=MaterialColors.BG_LIGHT, highlightthickness=0)
        scrollbar = ModernScrollbar(left_container, orient="vertical", command=canvas.yview)
        self.form_frame = tk.Frame(canvas, bg=MaterialColors.BG_LIGHT)
        
        self.form_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        
        def on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
        
        canvas.bind('<Configure>', on_canvas_configure)
        canvas_window = canvas.create_window((0, 0), window=self.form_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        row_counter = 0
        
        self.add_section("Tipo de Reparo", row_counter, icon="")
        row_counter += 1
        
        combo_frame = tk.Frame(self.form_frame, bg=MaterialColors.BG_LIGHT)
        combo_frame.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        
        self.report_type = ttk.Combobox(combo_frame, 
                                       values=["OPENED", "REOPENED", "VERIFIED"],
                                       state='readonly', 
                                       font=('Segoe UI', 11), 
                                       width=28,
                                       style='Modern.TCombobox')
        self.report_type.pack(fill=tk.X)
        self.report_type.current(0)
        self.report_type.bind('<<ComboboxSelected>>', self.on_type_change)
        row_counter += 1
        
        summary_card = self.add_section("Summary", row_counter, icon="")
        self.summary_widgets.append(summary_card)
        row_counter += 1
        
        summary_container = self.create_rounded_frame(self.form_frame, bg='white')
        summary_container.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        self.summary = tk.Entry(summary_container, font=('Segoe UI', 11), bg='white', 
                               relief=tk.FLAT, borderwidth=0)
        self.summary.pack(fill=tk.X, padx=12, pady=10)
        self.summary_widgets.append(summary_container)
        row_counter += 1
        
        self.add_section("Equipment Information", row_counter, icon="")
        row_counter += 1
        
        eq_frame = self.create_rounded_frame(self.form_frame, bg='white')
        eq_frame.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        
        eq_text_frame = tk.Frame(eq_frame, bg='white')
        eq_text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        self.equipment = tk.Text(eq_text_frame, font=('Segoe UI', 10), bg='white', height=8,
                                relief=tk.FLAT, borderwidth=0, wrap=tk.WORD, undo=True, maxundo=-1)
        eq_scroll = ModernScrollbar(eq_text_frame, orient="vertical", command=self.equipment.yview)
        self.equipment.configure(yscrollcommand=eq_scroll.set)
        
        self.equipment.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        eq_scroll.pack(side=tk.RIGHT, fill=tk.Y, padx=(0, 5), pady=10)
        row_counter += 1
        
        self.add_section("Descripcion", row_counter, icon="")
        row_counter += 1
        
        desc_frame = self.create_rounded_frame(self.form_frame, bg='white')
        desc_frame.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        
        desc_text_frame = tk.Frame(desc_frame, bg='white')
        desc_text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        self.description = tk.Text(desc_text_frame, font=('Segoe UI', 10), bg='white', height=8,
                                  relief=tk.FLAT, borderwidth=0, wrap=tk.WORD, undo=True, maxundo=-1)
        desc_scroll = ModernScrollbar(desc_text_frame, orient="vertical", command=self.description.yview)
        self.description.configure(yscrollcommand=desc_scroll.set)
        
        self.description.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        desc_scroll.pack(side=tk.RIGHT, fill=tk.Y, padx=(0, 5), pady=10)
        row_counter += 1
        
        # ========== LOGS DE CONSOLA ==========
        logs_label = self.add_section("Logs de Consola (Opcional)", row_counter, icon="")
        self.always_visible_widgets.append(logs_label)
        row_counter += 1
        
        logs_frame = self.create_rounded_frame(self.form_frame, bg='#2b2b2b')
        logs_frame.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        self.always_visible_widgets.append(logs_frame)
        
        logs_text_frame = tk.Frame(logs_frame, bg='#2b2b2b')
        logs_text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        self.console_logs = tk.Text(
            logs_text_frame, 
            font=('Consolas', 9),
            bg='#1e1e1e',
            fg='#00ff00',
            height=6,
            relief=tk.FLAT, 
            borderwidth=0, 
            wrap=tk.WORD, 
            undo=True, 
            maxundo=-1,
            insertbackground='#00ff00'
        )
        
        logs_scroll = ModernScrollbar(logs_text_frame, orient="vertical", command=self.console_logs.yview)
        self.console_logs.configure(yscrollcommand=logs_scroll.set)
        
        self.console_logs.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        logs_scroll.pack(side=tk.RIGHT, fill=tk.Y, padx=(0, 5), pady=10)
        
        placeholder_logs = "# Pega aquí los logs de consola (opcional)\n# No se traducirán ni corregirán\n# Ejemplo:\n# [ERROR] Connection timeout at 192.168.1.1\n# [INFO] Retry attempt 3/5..."
        self.console_logs.insert('1.0', placeholder_logs)
        self.console_logs.config(fg='#666666')
        
        def on_logs_focus_in(event):
            if self.console_logs.get('1.0', 'end-1c') == placeholder_logs:
                self.console_logs.delete('1.0', tk.END)
                self.console_logs.config(fg='#00ff00')
        
        def on_logs_focus_out(event):
            if not self.console_logs.get('1.0', 'end-1c').strip():
                self.console_logs.insert('1.0', placeholder_logs)
                self.console_logs.config(fg='#666666')
        
        self.console_logs.bind('<FocusIn>', on_logs_focus_in)
        self.console_logs.bind('<FocusOut>', on_logs_focus_out)
        row_counter += 1
        # =====================================

        proc_card = self.add_section("Procedimiento", row_counter, icon="")
        self.procedure_widgets.append(proc_card)
        row_counter += 1
        
        proc_frame = self.create_rounded_frame(self.form_frame, bg='white')
        proc_frame.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 15))
        
        proc_text_frame = tk.Frame(proc_frame, bg='white')
        proc_text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        self.procedure = AutoNumberedText(proc_text_frame, font=('Segoe UI', 10), bg='white', height=6,
                                         relief=tk.FLAT, borderwidth=0, wrap=tk.WORD)
        proc_scroll = ModernScrollbar(proc_text_frame, orient="vertical", command=self.procedure.yview)
        self.procedure.configure(yscrollcommand=proc_scroll.set)
        
        self.procedure.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        proc_scroll.pack(side=tk.RIGHT, fill=tk.Y, padx=(0, 5), pady=10)
        
        self.procedure.insert('1.0', '1. ')
        self.procedure_widgets.append(proc_frame)
        row_counter += 1
        
        btn_reset_container = tk.Frame(self.form_frame, bg=MaterialColors.BG_LIGHT)
        btn_reset_container.grid(row=row_counter, column=0, sticky='w', padx=30, pady=(0, 25))
        
        proc_btn = RoundedButton(btn_reset_container, text="Reiniciar", 
                                command=self.reset_proc,
                                bg_color=MaterialColors.TEXT_SECONDARY,
                                hover_color='#525252',
                                width=120, height=36, corner_radius=6)
        proc_btn.pack()
        self.procedure_widgets.append(btn_reset_container)
        row_counter += 1
        
        exp_card = self.add_section("Resultado Esperado", row_counter, icon="")
        self.expected_widgets.append(exp_card)
        row_counter += 1
        
        exp_frame = self.create_rounded_frame(self.form_frame, bg='white')
        exp_frame.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        
        exp_text_frame = tk.Frame(exp_frame, bg='white')
        exp_text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        self.expected = tk.Text(exp_text_frame, font=('Segoe UI', 10), bg='white', height=4,
                               relief=tk.FLAT, borderwidth=0, wrap=tk.WORD, undo=True, maxundo=-1)
        exp_scroll = ModernScrollbar(exp_text_frame, orient="vertical", command=self.expected.yview)
        self.expected.configure(yscrollcommand=exp_scroll.set)
        
        self.expected.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        exp_scroll.pack(side=tk.RIGHT, fill=tk.Y, padx=(0, 5), pady=10)
        self.expected_widgets.append(exp_frame)
        row_counter += 1
        
        self.add_section("Attachments", row_counter, icon="")
        row_counter += 1
        
        att_container = self.create_rounded_frame(self.form_frame, bg='white')
        att_container.grid(row=row_counter, column=0, sticky='ew', padx=30, pady=(0, 25))
        self.attachments = tk.Entry(att_container, font=('Segoe UI', 11), bg='white', 
                                    relief=tk.FLAT, borderwidth=0)
        self.attachments.pack(fill=tk.X, padx=12, pady=10)
        row_counter += 1
        
        btn_frame = tk.Frame(self.form_frame, bg=MaterialColors.BG_LIGHT)
        btn_frame.grid(row=row_counter, column=0, pady=35)
        
        btn_clear = RoundedButton(btn_frame, text="Limpiar", 
                                 command=self.clear_form,
                                 bg_color=MaterialColors.TEXT_SECONDARY,
                                 hover_color='#525252',
                                 width=140, height=45)
        btn_clear.pack(side=tk.LEFT, padx=8)
        
        btn_generate = RoundedButton(btn_frame, text="Generar", 
                                    command=self.generate,
                                    bg_color=MaterialColors.SUCCESS,
                                    hover_color=MaterialColors.SUCCESS_HOVER,
                                    font=('Segoe UI', 12, 'bold'),
                                    width=180, height=48)
        btn_generate.pack(side=tk.LEFT, padx=8)
        
        self.form_frame.columnconfigure(0, weight=1)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        def smart_scroll(event):
            widget = event.widget
            
            if isinstance(widget, tk.Text):
                scroll_position = widget.yview()
                delta = int(-1 * (event.delta / 120))
                
                if delta < 0:
                    if scroll_position[0] <= 0.0:
                        canvas.yview_scroll(delta, "units")
                        return "break"
                    else:
                        widget.yview_scroll(delta, "units")
                        return "break"
                else:
                    if scroll_position[1] >= 1.0:
                        canvas.yview_scroll(delta, "units")
                        return "break"
                    else:
                        widget.yview_scroll(delta, "units")
                        return "break"
            else:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
                return "break"
        
        def bind_smart_scroll_recursive(widget):
            widget.bind("<MouseWheel>", smart_scroll, add="+")
            for child in widget.winfo_children():
                bind_smart_scroll_recursive(child)
        
        bind_smart_scroll_recursive(self.form_frame)
        canvas.bind("<MouseWheel>", smart_scroll, add="+")
        
        right_container = tk.Frame(main_paned, bg=MaterialColors.BG_LIGHT)
        
        preview_card = self.create_rounded_frame(right_container, bg='white')
        preview_card.pack(fill=tk.BOTH, expand=True)
        
        preview_header = tk.Frame(preview_card, bg=MaterialColors.SUCCESS, height=55)
        preview_header.pack(fill=tk.X)
        preview_header.pack_propagate(False)
        
        tk.Label(preview_header, text="Vista Previa", font=('Segoe UI', 15, 'bold'),
                bg=MaterialColors.SUCCESS, fg='white').pack(pady=12)
        
        preview_container = tk.Frame(preview_card, bg='white')
        preview_container.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        preview_text_frame = tk.Frame(preview_container, bg='white')
        preview_text_frame.pack(fill=tk.BOTH, expand=True)
        
        self.preview = tk.Text(preview_text_frame, font=('Consolas', 10),
            bg='white', fg=MaterialColors.TEXT_PRIMARY, wrap=tk.WORD, 
            padx=18, pady=18, relief=tk.FLAT)
        preview_scroll = ModernScrollbar(preview_text_frame, orient="vertical", command=self.preview.yview)
        self.preview.configure(yscrollcommand=preview_scroll.set)
        
        self.preview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        preview_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.preview.insert('1.0', "\n\n    Vista Previa del Informe\n\n    "
                           "Completa el formulario y haz clic en Generar\n\n    "
                           "Se traducirá automáticamente al inglés\n    "
                           "Se corregirán errores gramaticales\n\n ")
        
        btn_preview = tk.Frame(preview_card, bg='white')
        btn_preview.pack(fill=tk.X, padx=15, pady=15)
        
        left_btns = tk.Frame(btn_preview, bg='white')
        left_btns.pack(side=tk.LEFT)
        
        btn_copy = RoundedButton(left_btns, text="Copiar", 
                                command=self.copy_preview,
                                bg_color=MaterialColors.PRIMARY,
                                hover_color=MaterialColors.PRIMARY_HOVER,
                                width=110, height=38)
        btn_copy.pack(side=tk.LEFT, padx=(0, 5))
        
        btn_export = RoundedButton(left_btns, text="Exportar", 
                                  command=self.export_word,
                                  bg_color=MaterialColors.SUCCESS,
                                  hover_color=MaterialColors.SUCCESS_HOVER,
                                  width=120, height=38)
        btn_export.pack(side=tk.LEFT, padx=(5, 0))
        
        right_btns = tk.Frame(btn_preview, bg='white')
        right_btns.pack(side=tk.RIGHT)
        
        btn_clear_prev = RoundedButton(right_btns, text="Limpiar", 
                                      command=self.clear_preview,
                                      bg_color=MaterialColors.TEXT_SECONDARY,
                                      hover_color='#525252',
                                      width=110, height=38)
        btn_clear_prev.pack()
        
        main_paned.add(left_container, minsize=600)
        main_paned.add(right_container, minsize=400)

        # ========== FOOTER CON CRÉDITOS ==========

        footer = tk.Frame(self.root, bg=MaterialColors.BG_LIGHT, height=35)
        footer.pack(fill=tk.X, side=tk.BOTTOM)
        footer.pack_propagate(False)

        tk.Label(
            footer, 
            text="© 2025 Luis Miguel Acosta & Coral Burgos",
            font=('Segoe UI', 10),
            bg=MaterialColors.BG_LIGHT,
            fg='#999999',
            anchor='center'
        ).pack(expand=True, fill=tk.BOTH)
    
    def setup_tab_order(self):
        widgets_order = [
            self.report_type,
            self.summary,
            self.equipment,
            self.description,
            self.console_logs,
            self.procedure,
            self.expected,
            self.attachments
        ]
        
        for i, widget in enumerate(widgets_order):
            if i < len(widgets_order) - 1:
                next_widget = widgets_order[i+1]
                
                if isinstance(widget, (ttk.Combobox, tk.Entry)):
                    widget.bind('<Tab>', lambda e, nw=next_widget: self.focus_next(nw, e))
                    widget.bind('<Shift-Tab>', lambda e, pw=widgets_order[i-1] if i > 0 else None: 
                               self.focus_previous(pw, e))
                elif isinstance(widget, tk.Text):
                    widget.bind('<Tab>', lambda e, nw=next_widget: self.focus_next(nw, e))
                    widget.bind('<Shift-Tab>', lambda e, pw=widgets_order[i-1] if i > 0 else None: 
                               self.focus_previous(pw, e))
        
        last_widget = widgets_order[-1]
        first_widget = widgets_order[0]
        if isinstance(last_widget, (ttk.Combobox, tk.Entry)):
            last_widget.bind('<Tab>', lambda e, fw=first_widget: self.focus_next(fw, e))
        elif isinstance(last_widget, tk.Text):
            last_widget.bind('<Tab>', lambda e, fw=first_widget: self.focus_next(fw, e))
    
    def focus_next(self, widget, event):
        widget.focus_set()
        return 'break'
    
    def focus_previous(self, widget, event):
        if widget:
            widget.focus_set()
        return 'break'
    
    def setup_keyboard_shortcuts(self):
        self.root.bind('<Control-s>', lambda e: (self.generate(), 'break'))
        self.root.bind('<Control-n>', lambda e: (self.clear_form(), 'break'))
        self.root.bind('<Control-e>', lambda e: (self.export_word(), 'break'))
        self.root.bind('<Control-q>', lambda e: self.root.quit())
        
        def clear_current_field(event):
            widget = self.root.focus_get()
            if isinstance(widget, tk.Entry):
                widget.delete(0, tk.END)
            elif isinstance(widget, tk.Text):
                widget.delete('1.0', tk.END)
                if widget == self.procedure:
                    self.procedure.insert('1.0', '1. ')
            return 'break'
        
        self.root.bind('<Escape>', clear_current_field)
    
    def add_section(self, title, row, icon=""):
        card = tk.Frame(self.form_frame, bg=MaterialColors.BG_LIGHT, relief=tk.FLAT)
        card.grid(row=row, column=0, sticky='ew', padx=20, pady=(15, 8))
        
        label_text = f"{icon} {title}" if icon else title
        tk.Label(card, text=label_text, font=('Segoe UI', 11, 'bold'),
                bg=MaterialColors.BG_LIGHT, fg=MaterialColors.PRIMARY,
                anchor='w').pack(fill=tk.X, pady=(0, 0))
        
        return card
    
    def on_type_change(self, event=None):
        report_type = self.report_type.get()
        
        for widget in self.summary_widgets:
            widget.grid_remove()
        for widget in self.procedure_widgets:
            widget.grid_remove()
        for widget in self.expected_widgets:
            widget.grid_remove()
        
        if report_type == "OPENED":
            for widget in self.summary_widgets:
                widget.grid()
            for widget in self.procedure_widgets:
                widget.grid()
            for widget in self.expected_widgets:
                widget.grid()
        elif report_type == "REOPENED":
            for widget in self.procedure_widgets:
                widget.grid()
        
        for widget in self.always_visible_widgets:
            widget.grid()
    
    def reset_proc(self):
        self.procedure.delete('1.0', tk.END)
        self.procedure.insert('1.0', '1. ')
    
    def clear_form(self):
        self.summary.delete(0, tk.END)
        self.equipment.delete('1.0', tk.END)
        self.description.delete('1.0', tk.END)
        
        placeholder_logs = "# Pega aquí los logs de consola (opcional)\n# No se traducirán ni corregirán\n# Ejemplo:\n# [ERROR] Connection timeout at 192.168.1.1\n# [INFO] Retry attempt 3/5..."
        self.console_logs.delete('1.0', tk.END)
        self.console_logs.insert('1.0', placeholder_logs)
        self.console_logs.config(fg='#666666')
        
        self.procedure.delete('1.0', tk.END)
        self.procedure.insert('1.0', '1. ')
        self.expected.delete('1.0', tk.END)
        self.attachments.delete(0, tk.END)
        self.report_type.current(0)
        self.on_type_change()
    
    def copy_preview(self):
        content = self.preview.get('1.0', 'end-1c')
        if not content.strip() or "Vista Previa" in content:
            messagebox.showwarning("Advertencia", "Genera primero un informe")
            return
        self.root.clipboard_clear()
        self.root.clipboard_append(content)
        messagebox.showinfo("Copiado", "Informe copiado al portapapeles")
    
    def clear_preview(self):
        self.preview.delete('1.0', tk.END)
        self.preview.insert('1.0', "\n\n    Vista Previa del Informe\n\n    "
                           "Completa el formulario y haz clic en Generar\n\n    "
                           "Se traducirá automáticamente al inglés\n    "
                           "Se corregirán errores gramaticales\n\n ")
    
    def export_word(self):
        content = self.preview.get('1.0', 'end-1c').strip()
        if not content or "Vista Previa" in content:
            messagebox.showwarning("Advertencia", "Genera primero un informe")
            return
        
        try:
            doc = Document()
            in_console_section = False
            
            for line in content.split('\n'):
                if line.strip():
                    if line.strip() in ['OPENED', 'REOPENED', 'VERIFIED']:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                        run.font.size = Pt(14)
                        if 'REOPENED' in line:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                        elif 'VERIFIED' in line:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                        in_console_section = False
                    elif line.strip().startswith('[') and line.strip().endswith(']:'):
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                        
                        if '[Console Logs]:' in line:
                            in_console_section = True
                        else:
                            in_console_section = False
                    else:
                        p = doc.add_paragraph()
                        
                        # Aplicar color especial para palabras clave en el texto
                        if 'VERIFIED' in line:
                            parts = line.split('VERIFIED')
                            for i, part in enumerate(parts):
                                if i > 0:
                                    run = p.add_run('VERIFIED')
                                    run.font.color.rgb = RGBColor(16, 124, 16)
                                    run.bold = True
                                if part:
                                    run = p.add_run(part)
                        elif 'REOPENED' in line:
                            parts = line.split('REOPENED')
                            for i, part in enumerate(parts):
                                if i > 0:
                                    run = p.add_run('REOPENED')
                                    run.font.color.rgb = RGBColor(204, 0, 0)
                                    run.bold = True
                                if part:
                                    run = p.add_run(part)
                        else:
                            run = p.add_run(line)
                        
                        # Aplicar formato de consola si estamos en esa sección
                        if in_console_section:
                            for run in p.runs:
                                run.font.name = 'Consolas'
                                run.font.size = Pt(9)
                                if run.font.color.rgb != RGBColor(16, 124, 16) and run.font.color.rgb != RGBColor(204, 0, 0):
                                    run.font.color.rgb = RGBColor(0, 102, 0)
                else:
                    doc.add_paragraph()
            
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word", "*.docx")],
                initialfile=f"Repair_{self.report_type.get()}_{ts}.docx"
            )
            if path:
                doc.save(path)
                messagebox.showinfo("Exportado", f"Documento guardado en:\n{path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def generate(self):
        self.root.focus_set()
        self.root.update_idletasks()
        
        rt = self.report_type.get()
        
        if rt == "OPENED":
            if not self.summary.get().strip():
                messagebox.showerror("Error", "Summary obligatorio para OPENED")
                return
            proc = self.procedure.get_numbered_text()
            if not proc.strip() or proc.strip() == "1.":
                messagebox.showerror("Error", "Procedimiento obligatorio para OPENED")
                return
            if not self.expected.get('1.0', tk.END).strip():
                messagebox.showerror("Error", "Resultado Esperado obligatorio para OPENED")
                return
        
        if not self.equipment.get('1.0', tk.END).strip():
            messagebox.showerror("Error", "Equipment Information requerido")
            return
        if not self.description.get('1.0', tk.END).strip():
            messagebox.showerror("Error", "Descripcion requerida")
            return
        
        try:
            self.preview.delete('1.0', tk.END)
            
            # Configurar tags de formato
            self.preview.tag_config("verified_word", foreground='#107C10', font=('Consolas', 10, 'bold'))
            self.preview.tag_config("reopened_word", foreground='#CC0000', font=('Consolas', 10, 'bold'))
            self.preview.tag_config("console_logs", font=('Consolas', 9), foreground='#006600', background='#f0f0f0')
            
            if rt == "VERIFIED":
                self.preview.insert(tk.END, "Traduciendo...\n")
                self.preview.see(tk.END)
                self.root.update()
                
                eq = translate_equipment_info(self.equipment.get('1.0', tk.END).strip())
                self.preview.delete("end-2l", "end-1l")
                self.preview.insert(tk.END, "[Equipment information]:\n")
                self.preview.insert(tk.END, f"{eq}\n\n")
                self.root.update()
                
                # Insertar "VERIFIED" con color verde
                verified_text = "The problem is VERIFIED in this version\n\n"
                start = self.preview.index("end-1c")
                self.preview.insert(tk.END, verified_text)
                
                # Aplicar tag solo a la palabra VERIFIED
                verified_start = f"{start} + {verified_text.index('VERIFIED')}c"
                verified_end = f"{verified_start} + 8c"
                self.preview.tag_add("verified_word", verified_start, verified_end)
                self.root.update()
                
                self.preview.insert(tk.END, "Traduciendo descripcion...\n")
                self.preview.see(tk.END)
                self.root.update()
                
                desc = translate_and_correct(self.description.get('1.0', tk.END).strip())
                self.preview.delete("end-2l", "end-1l")
                self.preview.insert(tk.END, f"{desc}\n\n")
                self.root.update()
                
                logs = self.console_logs.get('1.0', tk.END).strip()
                
                if logs and not logs.startswith("#"):
                    self.preview.insert(tk.END, "[Console Logs]:\n")
                    start_idx = self.preview.index("end-1c")
                    self.preview.insert(tk.END, f"{logs}\n\n")
                    end_idx = self.preview.index("end-1c")
                    
                    self.preview.tag_add("console_logs", start_idx, end_idx)
                    self.root.update()
                
                att = self.attachments.get().strip()
                if att:
                    self.preview.insert(tk.END, "[Attachments]:\n")
                    self.preview.insert(tk.END, f"{att}\n")
                
                messagebox.showinfo("Listo", "Informe generado correctamente")
                return
            
            self.preview.insert(tk.END, f"{rt}\n")
            if rt == "REOPENED":
                reopened_text = "The problem continues, REOPENED in this version.\n\n"
                start = self.preview.index("end-1c")
                self.preview.insert(tk.END, reopened_text)
                
                # Aplicar tag solo a la palabra REOPENED
                reopened_start = f"{start} + {reopened_text.index('REOPENED')}c"
                reopened_end = f"{reopened_start} + 8c"
                self.preview.tag_add("reopened_word", reopened_start, reopened_end)
            else:
                self.preview.insert(tk.END, "\n")
            self.root.update()
            
            if rt == "OPENED":
                self.preview.insert(tk.END, "Traduciendo...\n")
                self.preview.see(tk.END)
                self.root.update()
                
                summ = translate_and_correct(self.summary.get())
                self.preview.delete("end-2l", "end-1l")
                self.preview.insert(tk.END, f"Summary: {summ}\n\n")
                self.root.update()
            
            self.preview.insert(tk.END, "Traduciendo...\n")
            self.preview.see(tk.END)
            self.root.update()
            
            eq = translate_equipment_info(self.equipment.get('1.0', tk.END).strip())
            self.preview.delete("end-2l", "end-1l")
            self.preview.insert(tk.END, "[Equipment information]:\n\n")
            self.preview.insert(tk.END, f"{eq}\n\n")
            self.root.update()
            
            self.preview.insert(tk.END, "Traduciendo...\n")
            self.preview.see(tk.END)
            self.root.update()
            
            fault = translate_and_correct(self.description.get('1.0', tk.END).strip())
            self.preview.delete("end-2l", "end-1l")
            self.preview.insert(tk.END, "[Fault]:\n")
            self.preview.insert(tk.END, f"{fault}\n\n")
            self.root.update()
            
            logs = self.console_logs.get('1.0', tk.END).strip()
            
            if logs and not logs.startswith("#"):
                self.preview.insert(tk.END, "[Console Logs]:\n")
                start_idx = self.preview.index("end-1c")
                self.preview.insert(tk.END, f"{logs}\n\n")
                end_idx = self.preview.index("end-1c")
                
                self.preview.tag_add("console_logs", start_idx, end_idx)
                self.root.update()
            
            proc = self.procedure.get_numbered_text()
            if proc.strip() and proc.strip() != "1.":
                self.preview.insert(tk.END, "Traduciendo...\n")
                self.preview.see(tk.END)
                self.root.update()
                
                proc_t = translate_and_correct(proc)
                self.preview.delete("end-2l", "end-1l")
                self.preview.insert(tk.END, "[Procedure]:\n")
                self.preview.insert(tk.END, f"{proc_t}\n\n")
                self.root.update()
            
            if rt == "OPENED":
                exp = self.expected.get('1.0', tk.END).strip()
                if exp:
                    self.preview.insert(tk.END, "Traduciendo...\n")
                    self.preview.see(tk.END)
                    self.root.update()
                    
                    exp_t = translate_and_correct(exp)
                    self.preview.delete("end-2l", "end-1l")
                    self.preview.insert(tk.END, "[Expected]:\n")
                    self.preview.insert(tk.END, f"{exp_t}\n\n")
                    self.root.update()
            
            att = self.attachments.get().strip()
            if att:
                self.preview.insert(tk.END, "[Attachments]:\n")
                self.preview.insert(tk.END, f"{att}\n")
            
            messagebox.showinfo("Listo", "Informe generado correctamente")
        
        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = RepairReportGenerator(root)
    root.mainloop()